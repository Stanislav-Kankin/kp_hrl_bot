from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .utils import set_montserrat_font, format_cost, format_count


def load_template(template_name, need_onprem=True):
    if not need_onprem and template_name == "template.docx":
        template_name = "template_no_onprem.docx"
    return Document(f"templates/{template_name}")


def fill_standard_template(doc, data):
    table = doc.tables[0]
    set_montserrat_font(doc)

    # Устанавливаем оптимальные ширины столбцов
    table.columns[0].width = Pt(90)  # Количество сотрудников
    table.columns[1].width = Pt(180)  # Тип лицензии
    table.columns[2].width = Pt(100)  # Стоимость
    table.columns[3].width = Pt(80)   # Кол-во
    table.columns[4].width = Pt(80)   # Срок
    table.columns[5].width = Pt(100)  # Итого

    def fill_cell(row, col, text, bold=False):
        cell = table.cell(row, col)
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = bold
                run.font.name = 'Montserrat'
                run.font.size = Pt(9)  # Уменьшаем размер шрифта

    # Заполняем данные
    employee_count = format_count(data["employee_license_count"])
    fill_cell(1, 0, employee_count, bold=True)

    # Базовая лицензия
    fill_cell(1, 1, "Базовая лицензия")
    fill_cell(1, 2, format_cost(data["base_license_cost"], with_ruble=True))
    fill_cell(1, 3, format_count(data["base_license_count"]))
    fill_cell(1, 4, "12")
    fill_cell(1, 5, format_cost(
        data["base_license_cost"] * data["base_license_count"],
        with_ruble=True))

    # Лицензия кадровика
    fill_cell(2, 1, "Лицензия кадровика")
    fill_cell(2, 2, format_cost(data["hr_license_cost"], with_ruble=True))
    fill_cell(2, 3, format_count(data["hr_license_count"]))
    fill_cell(2, 4, "12")
    fill_cell(2, 5, format_cost(
        data["hr_license_cost"] * data["hr_license_count"],
        with_ruble=True))

    # Лицензия сотрудника
    fill_cell(3, 1, "Лицензия Сотрудника")
    fill_cell(3, 2, format_cost(data["employee_license_cost"],
                                with_ruble=True))
    fill_cell(3, 3, employee_count)
    fill_cell(3, 4, "12")
    fill_cell(3, 5, format_cost(
        data["employee_license_cost"] * data["employee_license_count"],
        with_ruble=True))

    # On-prem (если нужно)
    if data.get("need_onprem", False):
        fill_cell(4, 1, "On-prem размещение")
        fill_cell(4, 2, format_cost(data["onprem_cost"], with_ruble=True))
        fill_cell(4, 3, format_count(data["onprem_count"]))
        fill_cell(4, 4, "12")
        fill_cell(4, 5, format_cost(
            data["onprem_cost"] * data["onprem_count"], with_ruble=True))

    # Итоговая сумма
    total = (
        data["base_license_cost"] * data["base_license_count"] +
        data["hr_license_cost"] * data["hr_license_count"] +
        data["employee_license_cost"] * data["employee_license_count"]
    )
    if data.get("need_onprem", False):
        total += data["onprem_cost"] * data["onprem_count"]

    total_row = 5 if data.get("need_onprem", False) else 4
    fill_cell(total_row, 5, format_cost(total, with_ruble=True), bold=True)

    # Объединяем ячейки в первом столбце
    if data.get("need_onprem", False):
        table.cell(1, 0).merge(table.cell(4, 0))
    else:
        table.cell(1, 0).merge(table.cell(3, 0))

    insert_footer_expiration(doc, data.get("kp_expiration", ""))


def fill_complex_template(doc, data):
    set_montserrat_font(doc)
    company_name = data.get('company_name', '')

    for paragraph in doc.paragraphs:
        if "Коммерческое предложение HRlink для компании" in paragraph.text:
            paragraph.clear()
            run1 = paragraph.add_run(
                "Коммерческое предложение HRlink для компании "
            )
            run1.bold = True
            run1.font.size = Pt(18)
            run2 = paragraph.add_run(f'"{company_name}"')
            run2.bold = True
            run2.font.color.rgb = RGBColor(0x44, 0x9D, 0xE6)
            run2.font.size = Pt(18)
            break

    if len(doc.tables) > 0:
        table = doc.tables[0]

        def fill_cell(row, col, text, bold=False):
            cell = table.cell(row, col)
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = bold
                    run.font.name = 'Montserrat'

        fill_cell(1, 2, format_cost(data["base_license_cost"],
                                    with_ruble=True))
        fill_cell(1, 3, format_count(data["base_license_count"]))
        fill_cell(1, 5, format_cost(data["base_license_cost"] * data[
            "base_license_count"], with_ruble=True))

        fill_cell(2, 2, format_cost(data["hr_license_cost"], with_ruble=True))
        fill_cell(2, 3, format_count(data["hr_license_count"]))
        fill_cell(2, 5, format_cost(data["hr_license_cost"] * data[
            "hr_license_count"], with_ruble=True))

        fill_cell(3, 2, format_cost(data["employee_license_cost"],
                                    with_ruble=True))
        fill_cell(3, 3, format_count(data["employee_license_count"]))
        fill_cell(3, 5, format_cost(data["employee_license_cost"] * data[
            "employee_license_count"], with_ruble=True))

        if data.get("need_onprem"):
            fill_cell(4, 2, format_cost(data["onprem_cost"], with_ruble=True))
            fill_cell(4, 3, format_count(data["onprem_count"]))
            fill_cell(4, 4, "12")
            fill_cell(4, 5, format_cost(data["onprem_cost"] * data[
                "onprem_count"], with_ruble=True))

        total = (data["base_license_cost"] * data["base_license_count"] +
                 data["hr_license_cost"] * data["hr_license_count"] +
                 data["employee_license_cost"] * data[
                     "employee_license_count"])
        if data.get("need_onprem"):
            total += data["onprem_cost"] * data["onprem_count"]

        total_row = 5 if data.get("need_onprem") else 4
        fill_cell(total_row, 5, format_cost(total, with_ruble=True), bold=True)

    insert_footer_expiration(doc, data.get("kp_expiration", ""))


def fill_marketing_template(doc, data):
    set_montserrat_font(doc)
    company_name = data.get('company_name', '')

    # Цвет и размер для заголовков
    title_color = RGBColor(0x44, 0x9D, 0xE6)
    title_font_size = Pt(13)

    # Обновляем название компании
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "HRlink для" in paragraph.text:
                        for run in paragraph.runs:
                            run.text = ""
                        run = paragraph.add_run(f"HRlink для {company_name}")
                        run.bold = True
                        run.font.size = Pt(15)
                        run.font.name = 'Montserrat'
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        break

    if len(doc.tables) <= 2:
        print("[!] Таблица с ценами не найдена.")
        return

    # Таблица с ценами
    table = doc.tables[2]
    rows_count = len(table.rows)
    cols_count = len(table.columns)

    def fill_cell(row, col, text, bold=False):
        if row >= rows_count or col >= cols_count:
            print(
                f"[!] Нет ячейки ({row}, {col}) в таблице {rows_count}x{cols_count}"
                )
            return
        cell = table.cell(row, col)
        cell.text = str(text)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = bold
                run.font.name = 'Montserrat'
                run.font.size = Pt(9)

    # Заполнение таблицы
    fill_cell(1, 1, format_cost(data["base_license_cost"], with_ruble=True))
    fill_cell(1, 2, format_count(data["base_license_count"]))
    fill_cell(1, 3, "12 мес.")
    fill_cell(1, 4, format_cost(
        data["base_license_cost"] * data[
            "base_license_count"], with_ruble=True))

    fill_cell(2, 1, format_cost(data["hr_license_cost"], with_ruble=True))
    fill_cell(2, 2, format_count(data["hr_license_count"]))
    fill_cell(2, 3, "12 мес.")
    fill_cell(2, 4, format_cost(
        data["hr_license_cost"] * data["hr_license_count"], with_ruble=True))

    fill_cell(3, 1, format_cost(data[
        "employee_license_cost"], with_ruble=True))
    fill_cell(3, 2, format_count(data["employee_license_count"]))
    fill_cell(3, 3, "12 мес.")
    fill_cell(3, 4, format_cost(
        data["employee_license_cost"] * data[
            "employee_license_count"], with_ruble=True))

    total = (
        data["base_license_cost"] * data["base_license_count"] +
        data["hr_license_cost"] * data["hr_license_count"] +
        data["employee_license_cost"] * data["employee_license_count"]
    )

    if data.get("need_onprem", False):
        fill_cell(4, 1, format_cost(data["onprem_cost"], with_ruble=True))
        fill_cell(4, 2, format_count(data["onprem_count"]))
        fill_cell(4, 3, "12 мес.")
        fill_cell(4, 4, format_cost(
            data["onprem_cost"] * data["onprem_count"], with_ruble=True))
        total += data["onprem_cost"] * data["onprem_count"]

    total_row = 5 if data.get("need_onprem", False) else 4
    fill_cell(total_row, 4, format_cost(total, with_ruble=True), bold=True)

    # Обновляем блок "На 1 лицензию Сотрудника"
    for paragraph in doc.paragraphs:
        if "На 1 лицензию Сотрудника" in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run("На 1 лицензию Сотрудника")
            run.bold = True
            run.font.color.rgb = title_color
            run.font.size = title_font_size

            content_run = paragraph.add_run("\n")
            content_run = paragraph.add_run(
                f"- {data.get('unep_count', 0)} УНЭП\n"
                f"- {data.get(
                    'sms_count', 0
                    )} СМС (на уведомление и подписание документа)"
            )
            content_run.font.size = Pt(10)
            content_run.font.color.rgb = RGBColor(0, 0, 0)
            break

    # Обновляем "Индивидуальные условия"
    conditions = data.get("custom_conditions", [])
    for i, paragraph in enumerate(doc.paragraphs):
        if "Индивидуальные условия" in paragraph.text:
            paragraph.clear()
            title_run = paragraph.add_run("Индивидуальные условия")
            title_run.bold = True
            title_run.font.color.rgb = title_color
            title_run.font.size = title_font_size

            # вставка условий как новые параграфы после текущего
            parent = paragraph._element.getparent()
            next_el = paragraph._element
            for condition in conditions or ["-"]:
                new_p = doc.add_paragraph(f"- {condition}")
                for run in new_p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                parent.insert(parent.index(next_el) + 1, new_p._element)
                next_el = new_p._element
            break

    insert_footer_expiration(doc, data.get("kp_expiration", ""))


def insert_footer_expiration(doc, date_text):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[
            0
            ] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = f"Коммерческое предложение действительно до {
            date_text} г."
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(10)
        run.font.name = 'Montserrat'
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)


def fill_396_template(doc, data, is_onprem=False):
    set_montserrat_font(doc)

    # Обрабатываем обе таблицы (обычный тариф и PRO)
    for table in doc.tables[:2]:
        def fill_cell(row, col, text, bold=False):
            cell = table.cell(row, col)
            cell.text = text
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = bold
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(10)

        employee_count = format_count(data["employee_license_count"])

        # Заполняем данные
        fill_cell(1, 0, employee_count, bold=True)

        # Базовая лицензия
        fill_cell(1, 1, "Базовая лицензия")
        fill_cell(1, 2, "15 000 ₽")
        fill_cell(1, 3, "1")
        fill_cell(1, 4, "12")
        fill_cell(1, 5, "15 000 ₽")

        # Лицензия кадровика
        fill_cell(2, 1, "Лицензия кадровика")
        fill_cell(2, 2, "15 000 ₽")
        fill_cell(2, 3, format_count(data["hr_license_count"]))
        fill_cell(2, 4, "12")
        fill_cell(2, 5, format_cost(15000 * data[
            "hr_license_count"], with_ruble=True))

        # Лицензия сотрудника
        fill_cell(3, 1, "Лицензия Сотрудника")
        fill_cell(3, 2, "396 ₽")
        fill_cell(3, 3, employee_count)
        fill_cell(3, 4, "12")
        fill_cell(3, 5, format_cost(396 * data[
            "employee_license_count"], with_ruble=True))

        # Определяем строки для разных типов таблиц
        if "PRO" in table.cell(4, 1).text:  # Это таблица PRO
            # Лицензия PRO
            fill_cell(4, 1, "Лицензия сотрудника PRO")
            fill_cell(4, 2, "144 ₽")
            fill_cell(4, 3, employee_count)
            fill_cell(4, 4, "12")
            fill_cell(4, 5, format_cost(144 * data[
                "employee_license_count"], with_ruble=True))
            onprem_row = 5
        else:
            onprem_row = 4

        # On-prem (если нужен)
        if is_onprem:
            fill_cell(onprem_row, 1, "On-prem размещение")
            fill_cell(onprem_row, 2, "400 000 ₽")
            fill_cell(onprem_row, 3, "1")
            fill_cell(onprem_row, 4, "12")
            fill_cell(onprem_row, 5, "400 000 ₽")
            total_row = onprem_row + 1
        else:
            total_row = onprem_row

        # SMS (фиксированная стоимость)
        fill_cell(total_row, 1, "SMS-сообщения*")
        fill_cell(total_row, 2, "4 ₽")
        fill_cell(total_row, 3, "1")
        fill_cell(total_row, 4, "12")
        fill_cell(total_row, 5, "4 ₽")

        # Итоговая сумма (без учета SMS)
        total = (
            15000 * 1 +  # Базовая лицензия
            15000 * data["hr_license_count"] +
            396 * data["employee_license_count"]
        )

        if "PRO" in table.cell(4, 1).text:
            total += 144 * data["employee_license_count"]

        if is_onprem:
            total += 400000

        fill_cell(total_row + 1, 5, format_cost(
            total, with_ruble=True), bold=True)

        # Объединяем ячейки в первом столбце
        if "PRO" in table.cell(4, 1).text:
            merge_to = 5 if is_onprem else 4
        else:
            merge_to = 4 if is_onprem else 3

        table.cell(1, 0).merge(table.cell(merge_to, 0))
