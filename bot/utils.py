import os
from collections import OrderedDict
from docx.shared import Pt

import subprocess
import shutil


file_id_mapping = OrderedDict()


def clean_input(value):
    try:
        # Округляем до целого числа
        return int(round(float(value.replace(',', '.').strip())))
    except ValueError:
        raise ValueError(f"Некорректное значение: {value}")


def format_cost(value, with_ruble=False):
    text = f"{int(round(float(value))):,}".replace(',', ' ')
    return f"{text} ₽" if with_ruble else text


def format_count(value):
    return f"{int(value)}"


def cleanup_kp_files():
    current_dir = os.getcwd()
    for filename in os.listdir(current_dir):
        if filename.startswith(
            "КП_") and (
                filename.endswith(".docx") or filename.endswith(".pdf")):
            os.remove(os.path.join(current_dir, filename))


def set_montserrat_font(doc):
    styles = doc.styles
    for style in styles:
        if style.type == 1:
            font = style.font
            font.name = 'Montserrat'
            font.size = Pt(10)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[
                        0
                        ] if paragraph.runs else paragraph.add_run()
                    run.font.name = 'Montserrat'
                    run.font.size = Pt(10)


def convert_to_pdf_libreoffice(input_path: str) -> str:
    """Конвертирует DOCX → PDF с помощью LibreOffice CLI"""
    output_dir = os.path.dirname(input_path)

    # Ищем путь к soffice/libreoffice
    soffice_path = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice_path or not os.path.exists(soffice_path):
        raise RuntimeError("LibreOffice (libreoffice или soffice) не найдена в системе.")

    # Подготавливаем окружение с безопасным PATH
    env = os.environ.copy()
    env["PATH"] += ":/usr/bin:/usr/local/bin"

    # Запускаем LibreOffice CLI
    result = subprocess.run(
        [soffice_path, "--headless", "--convert-to", "pdf", input_path, "--outdir", output_dir],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        env=env
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"Ошибка конвертации LibreOffice:\n"
            f"{result.stderr.decode(errors='ignore') or result.stdout.decode(errors='ignore')}"
        )

    pdf_path = os.path.splitext(input_path)[0] + ".pdf"
    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF файл не создан — возможно, формат исходника не поддерживается.")

    return pdf_path

