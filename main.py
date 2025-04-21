import os
import uuid
import threading
from collections import OrderedDict
from datetime import datetime, timedelta
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import InlineKeyboardButton, InlineKeyboardMarkup
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
from PyPDF2 import PdfReader, PdfWriter

# Загрузка токена из .env
load_dotenv()
TOKEN = os.getenv("BOT_TOKEN")

bot = Bot(token=TOKEN)
dp = Dispatcher()

# Состояния для FSM
class Form(StatesGroup):
    base_license_cost = State()
    base_license_count = State()
    hr_license_cost = State()
    hr_license_count = State()
    employee_license_cost = State()
    employee_license_count = State()
    need_onprem = State()
    onprem_cost = State()
    onprem_count = State()

# Словарь для хранения соответствия между уникальным идентификатором и file_id
file_id_mapping = OrderedDict()

# Словарь для хранения времени создания файлов
file_creation_times = OrderedDict()

# Функция для проверки и очистки данных
def clean_input(value):
    try:
        return float(value.replace(',', '.').strip())
    except ValueError:
        raise ValueError(f"Некорректное значение: {value}")

# Функция для форматирования стоимости
def format_cost(value):
    return f"{value:,.2f}".replace(',', ' ').replace('.', ',')

# Функция для форматирования количества
def format_count(value):
    return f"{int(value)}"

# Функция для удаления старых файлов
def cleanup_old_files():
    while True:
        current_time = datetime.now()
        files_to_delete = []
        for file_path, creation_time in file_creation_times.items():
            if current_time - creation_time > timedelta(minutes=10):  # Удаляем файлы старше 10 минут
                files_to_delete.append(file_path)

        for file_path in files_to_delete:
            if os.path.exists(file_path):
                os.remove(file_path)
            del file_creation_times[file_path]

        # Ограничиваем количество файлов до 5
        while len(file_creation_times) > 5:
            oldest_file = next(iter(file_creation_times))
            if os.path.exists(oldest_file):
                os.remove(oldest_file)
            del file_creation_times[oldest_file]

        threading.Event().wait(60)  # Проверяем каждую минуту

# Запускаем поток для очистки старых файлов
threading.Thread(target=cleanup_old_files, daemon=True).start()

# Обработчик команды /start
@dp.message(Command("start"))
async def start(message: types.Message):
    await message.answer(
        "Это бот для создания КП. Нажмите /kp для начала."
    )

# Обработчик команды /kp
@dp.message(Command("kp"))
async def start_kp(message: types.Message, state: FSMContext):
    await state.set_state(Form.base_license_cost)
    await message.answer("Введите стоимость Базовой лицензии (руб/год):")

# Обработчик стоимости Базовой лицензии
@dp.message(Form.base_license_cost)
async def process_base_license_cost(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(base_license_cost=value)
        await state.set_state(Form.base_license_count)
        await message.answer("Введите количество Базовых лицензий:")
    except ValueError as e:
        await message.answer(str(e))

# Обработчик количества Базовых лицензий
@dp.message(Form.base_license_count)
async def process_base_license_count(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(base_license_count=value)
        await state.set_state(Form.hr_license_cost)
        await message.answer("Введите стоимость лицензий кадровиков (руб/год):")
    except ValueError as e:
        await message.answer(str(e))

# Обработчик стоимости лицензий кадровиков
@dp.message(Form.hr_license_cost)
async def process_hr_license_cost(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(hr_license_cost=value)
        await state.set_state(Form.hr_license_count)
        await message.answer("Введите количество лицензий кадровиков:")
    except ValueError as e:
        await message.answer(str(e))

# Обработчик количества лицензий кадровиков
@dp.message(Form.hr_license_count)
async def process_hr_license_count(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(hr_license_count=value)
        await state.set_state(Form.employee_license_cost)
        await message.answer("Введите стоимость лицензии сотрудника (руб/год):")
    except ValueError as e:
        await message.answer(str(e))

# Обработчик стоимости лицензии сотрудника
@dp.message(Form.employee_license_cost)
async def process_employee_license_cost(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(employee_license_cost=value)
        await state.set_state(Form.employee_license_count)
        await message.answer("Введите количество лицензий сотрудника:")
    except ValueError as e:
        await message.answer(str(e))

# Обработчик количества лицензий сотрудника
@dp.message(Form.employee_license_count)
async def process_employee_license_count(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(employee_license_count=value)
        await state.set_state(Form.need_onprem)

        # Создаем инлайн-кнопки
        keyboard = InlineKeyboardMarkup(inline_keyboard=[
            [InlineKeyboardButton(text="Да", callback_data="onprem_yes")],
            [InlineKeyboardButton(text="Нет", callback_data="onprem_no")]
        ])
        await message.answer("Нужен ли on-prem?", reply_markup=keyboard)
    except ValueError as e:
        await message.answer(str(e))

# Обработчик ответа на вопрос про on-prem
@dp.callback_query(F.data.startswith("onprem_"))
async def process_onprem_choice(callback: types.CallbackQuery, state: FSMContext):
    choice = callback.data.split("_")[1]

    if choice == "yes":
        await state.update_data(need_onprem=True)
        await state.set_state(Form.onprem_cost)
        await callback.message.answer("Введите сумму on-prem (руб/год):")
    else:
        await state.update_data(need_onprem=False, onprem_cost=0, onprem_count=0)
        await generate_kp(callback.message, state)

    await callback.answer()

# Обработчик суммы on-prem
@dp.message(Form.onprem_cost)
async def process_onprem_cost(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(onprem_cost=value)
        await state.set_state(Form.onprem_count)
        await message.answer("Введите количество лицензий on-prem:")
    except ValueError as e:
        await message.answer(str(e))

# Обработчик количества лицензий on-prem
@dp.message(Form.onprem_count)
async def process_onprem_count(message: types.Message, state: FSMContext):
    try:
        value = clean_input(message.text)
        await state.update_data(onprem_count=value)
        await generate_kp(message, state)
    except ValueError as e:
        await message.answer(str(e))

# Генерация КП и отправка пользователю
async def generate_kp(message: types.Message, state: FSMContext):
    data = await state.get_data()

    # Загружаем шаблон
    doc = Document("template.docx")

    # Обновляем таблицу (пример для всех строк)
    table = doc.tables[0]  # Предполагаем, что таблица первая в документе

    # Заполняем данные (адаптируйте под структуру вашего шаблона)
    table.cell(1, 2).text = format_cost(data["base_license_cost"])  # Стоимость Базовой лицензии
    table.cell(1, 3).text = format_count(data["base_license_count"])  # Количество
    table.cell(1, 5).text = format_cost(data["base_license_cost"] * data["base_license_count"])  # Итого

    table.cell(2, 2).text = format_cost(data["hr_license_cost"])  # Стоимость лицензий кадровиков
    table.cell(2, 3).text = format_count(data["hr_license_count"])  # Количество
    table.cell(2, 5).text = format_cost(data["hr_license_cost"] * data["hr_license_count"])  # Итого

    table.cell(3, 2).text = format_cost(data["employee_license_cost"])  # Стоимость лицензии сотрудника
    table.cell(3, 3).text = format_count(data["employee_license_count"])  # Количество
    table.cell(3, 5).text = format_cost(data["employee_license_cost"] * data["employee_license_count"])  # Итого

    if data["need_onprem"]:
        table.cell(4, 2).text = format_cost(data["onprem_cost"])  # Стоимость on-prem
        table.cell(4, 3).text = format_count(data["onprem_count"])  # Количество
        table.cell(4, 4).text = "12"  # Срок, мес
        table.cell(4, 5).text = format_cost(data["onprem_cost"] * data["onprem_count"])  # Итого
    else:
        table.cell(4, 2).text = "-"  # Прочерк для стоимости on-prem
        table.cell(4, 3).text = "-"  # Прочерк для количества on-prem
        table.cell(4, 4).text = "-"  # Прочерк для срока, мес
        table.cell(4, 5).text = "-"  # Прочерк для итого on-prem

    # Вычисляем итоговую сумму
    total = (data["base_license_cost"] * data["base_license_count"] +
             data["hr_license_cost"] * data["hr_license_count"] +
             data["employee_license_cost"] * data["employee_license_count"])
    if data["need_onprem"]:
        total += data["onprem_cost"] * data["onprem_count"]

    # Заполняем строку "ИТОГО"
    table.cell(5, 5).text = format_cost(total)  # Итого

    # Сохраняем измененный документ
    kp_filename = f"КП_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(kp_filename)

    # Отправляем файл пользователю
    with open(kp_filename, 'rb') as file:
        doc_message = await message.answer_document(
            types.BufferedInputFile(file.read(), filename=kp_filename),
            caption="Ваше КП готово! Вы можете скачать его или конвертировать в PDF."
        )

    # Генерируем уникальный идентификатор для callback_data
    unique_id = str(uuid.uuid4())

    # Ограничиваем количество файлов до 5
    if len(file_id_mapping) >= 5:
        oldest_file_id = next(iter(file_id_mapping))
        del file_id_mapping[oldest_file_id]

    file_id_mapping[unique_id] = doc_message.document.file_id
    file_creation_times[kp_filename] = datetime.now()

    # Создаем инлайн-кнопку для конвертации в PDF
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Сделать PDF", callback_data=f"convert_to_pdf_{unique_id}")]
    ])
    await message.answer("Можно сделать из этого фала PDF формат.\nЖми кнопку ниже!", reply_markup=keyboard)

    # Очищаем состояние
    await state.clear()

# Обработчик конвертации в PDF
@dp.callback_query(F.data.startswith("convert_to_pdf_"))
async def convert_to_pdf(callback: types.CallbackQuery):
    unique_id = callback.data.split("_")[3]
    file_id = file_id_mapping.get(unique_id)

    if not file_id:
        await callback.message.answer("Файл не найден.")
        await callback.answer()
        return

    file_info = await bot.get_file(file_id)
    file_path = file_info.file_path

    # Скачиваем файл
    file = await bot.download_file(file_path)
    doc_bytes = BytesIO(file.read())

    # Конвертируем DOCX в PDF
    pdf_bytes = BytesIO()
    doc = Document(doc_bytes)
    doc.save(pdf_bytes)
    pdf_bytes.seek(0)

    # Отправляем PDF пользователю
    pdf_filename = f"КП_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    await callback.message.answer_document(
        types.BufferedInputFile(pdf_bytes.read(), filename=pdf_filename),
        caption="Ваше КП в формате PDF."
    )

    await callback.answer()

# Запуск бота
if __name__ == "__main__":
    dp.run_polling(bot)
